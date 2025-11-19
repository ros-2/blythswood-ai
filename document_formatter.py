"""
Ross AI Grant Writer - Document Formatting
Clean Word/PDF output with no markdown artifacts
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER
import re
import io

def clean_markdown(text):
    """Remove all markdown formatting."""
    # Remove headers (##, ###, etc)
    text = re.sub(r'#{1,6}\s+', '', text)
    # Remove bold (**text** or __text__)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    # Remove italic (*text* or _text_)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    # Remove bullet points
    text = re.sub(r'^\s*[-*+]\s+', '', text, flags=re.MULTILINE)
    # Clean up extra whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def parse_sections(text):
    """
    Parse text into sections based on headers.
    Returns list of (title, content) tuples.
    """
    sections = []
    
    # Split by ## headers
    parts = re.split(r'##\s+(.+?)\n', text)
    
    if len(parts) > 1:
        # First part is intro (if exists)
        if parts[0].strip():
            sections.append(("Introduction", clean_markdown(parts[0])))
        
        # Process pairs of (title, content)
        for i in range(1, len(parts), 2):
            if i+1 < len(parts):
                title = parts[i].strip()
                content = clean_markdown(parts[i+1])
                sections.append((title, content))
    else:
        # No sections found, treat as single block
        sections.append(("Grant Application", clean_markdown(text)))
    
    return sections

def create_word_document(text, charity_name="Charity"):
    """
    Create properly formatted Word document.
    No markdown artifacts.
    """
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # Title
    title = doc.add_heading('Grant Application', 0)
    title_format = title.runs[0]
    title_format.font.name = 'Arial'
    title_format.font.size = Pt(18)
    title_format.font.bold = True
    
    # Subtitle
    subtitle = doc.add_paragraph(f'{charity_name}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.name = 'Arial'
    subtitle_format.font.size = Pt(12)
    subtitle_format.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()  # Spacing
    
    # Parse and add sections
    sections_list = parse_sections(text)
    
    for title, content in sections_list:
        # Section header
        heading = doc.add_heading(title, level=1)
        heading_format = heading.runs[0]
        heading_format.font.name = 'Arial'
        heading_format.font.size = Pt(14)
        heading_format.font.bold = True
        heading_format.font.color.rgb = RGBColor(0, 0, 0)
        
        # Section content - split into paragraphs
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para_format = para.runs[0] if para.runs else None
                if para_format:
                    para_format.font.name = 'Arial'
                    para_format.font.size = Pt(11)
        
        doc.add_paragraph()  # Spacing between sections
    
    # Save to bytes
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    
    return bio.getvalue()

def create_pdf_document(text, charity_name="Charity"):
    """
    Create properly formatted PDF.
    No markdown artifacts.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        topMargin=1*inch,
        bottomMargin=1*inch,
        leftMargin=1*inch,
        rightMargin=1*inch
    )
    
    # Styles
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=18,
        textColor=RGBColor(0, 0, 0),
        alignment=TA_CENTER,
        spaceAfter=6
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=12,
        textColor=RGBColor(102, 102, 102),
        alignment=TA_CENTER,
        spaceAfter=20
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=14,
        textColor=RGBColor(0, 0, 0),
        spaceAfter=12,
        spaceBefore=12
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=11,
        textColor=RGBColor(0, 0, 0),
        alignment=TA_LEFT,
        spaceAfter=12
    )
    
    # Build document
    story = []
    
    # Title
    story.append(Paragraph("Grant Application", title_style))
    story.append(Paragraph(charity_name, subtitle_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Parse and add sections
    sections_list = parse_sections(text)
    
    for section_title, content in sections_list:
        # Section heading
        story.append(Paragraph(section_title, heading_style))
        
        # Section content
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                # Escape special characters for ReportLab
                para_text = para_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(para_text.strip(), body_style))
        
        story.append(Spacer(1, 0.2*inch))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    
    return buffer.getvalue()

def create_text_file(text):
    """
    Create clean text file (just remove markdown).
    """
    sections_list = parse_sections(text)
    
    output = "GRANT APPLICATION\n"
    output += "=" * 50 + "\n\n"
    
    for title, content in sections_list:
        output += title.upper() + "\n"
        output += "-" * len(title) + "\n\n"
        output += content + "\n\n"
    
    return output
